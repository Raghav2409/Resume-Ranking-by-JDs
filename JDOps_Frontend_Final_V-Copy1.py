import streamlit as st
import boto3
import json
import pandas as pd
import os
from docx import Document
import re
from collections import Counter
import numpy as np
import plotly.graph_objects as go
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import nltk
import datetime
import time
from typing import Dict, List, Any, Optional
import uuid
from jdoptim_logger import JDOptimLogger

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    
def read_job_description(file_path):
    """Read job description from either .txt or .docx file"""
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format")

# Function to save enhanced JD
def save_enhanced_jd(content, filename, format_type):
    if format_type == 'docx':
        doc = Document()
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(filename)
        return True
    return False

class JobDescriptionAnalyzer:
    def __init__(self):
        self.categories = {
            'Technical Skills': ['python', 'java', 'sql', 'aws', 'cloud', 'docker', 'kubernetes', 'api', 'database', 
                               'git', 'linux', 'agile', 'devops', 'ml', 'ai', 'analytics'],
            'Soft Skills': ['communication', 'leadership', 'teamwork', 'collaboration', 'problem-solving', 'analytical', 
                           'initiative', 'organizational', 'time management', 'interpersonal'],
            'Experience Level': ['year', 'years', 'senior', 'junior', 'mid-level', 'lead', 'manager', 'experience'],
            'Education': ['degree', 'bachelor', 'master', 'phd', 'certification', 'education'],
            'Tools & Technologies': ['jira', 'confluence', 'slack', 'github', 'gitlab', 'azure', 'jenkins', 'terraform'],
            'Domain Knowledge': ['finance', 'healthcare', 'retail', 'banking', 'insurance', 'technology', 'manufacturing']
        }

    def analyze_text(self, text):
        text = text.lower()
        scores = {}
        
        for category, keywords in self.categories.items():
            category_score = 0
            for keyword in keywords:
                count = len(re.findall(r'\b' + keyword + r'\b', text))
                category_score += count
            max_possible = len(keywords)
            scores[category] = min(category_score / max_possible, 1.0)
            
        return scores

class JobDescriptionAgent:
    def __init__(self, model_id, max_tokens=1500, temperature=0.7):
        self.model_id = model_id
        self.max_tokens = max_tokens
        self.temperature = temperature
        
        # Get AWS credentials from environment variables or use a proper credential provider
        # SECURITY: Replace hardcoded credentials with proper credential management
        self.client = boto3.client(
                service_name='bedrock-runtime',
                aws_access_key_id='',
                aws_secret_access_key='',
                region_name='us-east-1',
            )

    def generate_initial_descriptions(self, job_description):
        """Generate detailed and structured job descriptions based on the given job description."""
        prompt = (
            "Using the job description provided, create three distinct enhanced versions. Each should expand on these sections:\n\n"
            "1. Brief overview of the role\n"
            "2. Key responsibilities\n"
            "3. Required skills\n"
            "4. Preferred skills\n"
            "5. Required experience\n"
            "6. Preferred experience\n"
            "7. Required tools\n\n"
            "Format your response exactly as follows:\n"
            "VERSION 1:\n[First enhanced version]\n\n"
            "VERSION 2:\n[Second enhanced version]\n\n"
            "VERSION 3:\n[Third enhanced version]\n\n"
            "Make each version unique while maintaining accuracy and relevance.\n\n"
            f"Original Job Description:\n{job_description}"
        )

        native_request = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": self.max_tokens,
            "temperature": self.temperature,
            "messages": [{"role": "user", "content": prompt}],
        }
        
        response = self.client.invoke_model(
            modelId=self.model_id,
            body=json.dumps(native_request),
            contentType="application/json",
        )
        response_body = response['body'].read().decode("utf-8")
        model_response = json.loads(response_body)

        if "content" in model_response and isinstance(model_response["content"], list):
            full_text = model_response["content"][0]["text"].strip()
            
            # More robust splitting pattern
            parts = re.split(r'VERSION \d+:', full_text)
            if len(parts) >= 4:  # The first part is empty or intro text
                descriptions = [part.strip() for part in parts[1:4]]
                return descriptions
            else:
                # Fallback parsing method
                descriptions = []
                version_pattern = re.compile(r'VERSION (\d+):(.*?)(?=VERSION \d+:|$)', re.DOTALL)
                matches = version_pattern.findall(full_text)
                for _, content in matches[:3]:
                    descriptions.append(content.strip())
                
                if len(descriptions) == 3:
                    return descriptions
        
        # If we failed to parse properly, generate simpler versions
        return [
            f"Enhanced Version 1 of the job description:\n{job_description}",
            f"Enhanced Version 2 of the job description:\n{job_description}",
            f"Enhanced Version 3 of the job description:\n{job_description}"
        ]

    def generate_final_description(self, selected_description, feedback_history):
        """
        Generate enhanced description incorporating feedback history
        
        Args:
            selected_description (str): The base description to enhance
            feedback_history (list): List of previous feedback items
        """
        # Construct prompt with feedback history
        feedback_context = "\n".join([
            f"Previous Feedback {i+1}: {feedback}" 
            for i, feedback in enumerate(feedback_history[:-1])
        ])
        
        current_feedback = feedback_history[-1] if feedback_history else ""
        
        prompt = (
            "Using the selected job description and incorporating all previous feedback plus new feedback, "
            "create an enhanced version.\n\n"
            f"Selected Job Description:\n{selected_description}\n\n"
        )
        
        if feedback_context:
            prompt += f"Previous Feedback Applied:\n{feedback_context}\n\n"
        
        prompt += (
            f"New Feedback to Implement:\n{current_feedback}\n\n"
            "Enhance while maintaining this structure:\n"
            "1. Brief overview of the role\n"
            "2. Key responsibilities\n"
            "3. Required skills\n"
            "4. Preferred skills\n"
            "5. Required experience\n"
            "6. Preferred experience\n"
            "7. Required tools\n\n"
            "Important: Ensure all previous feedback improvements are preserved while implementing the new feedback."
        )

        native_request = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": self.max_tokens,
            "temperature": self.temperature,
            "messages": [{"role": "user", "content": prompt}],
        }
        
        response = self.client.invoke_model(
            modelId=self.model_id,
            body=json.dumps(native_request),
            contentType="application/json",
        )
        response_body = response['body'].read().decode("utf-8")
        model_response = json.loads(response_body)

        if "content" in model_response and isinstance(model_response["content"], list):
            return model_response["content"][0]["text"].strip()
        else:
            raise ValueError("Unexpected response format")

def create_multi_radar_chart(scores_dict):
    """Create a radar chart comparing multiple job descriptions"""
    categories = list(next(iter(scores_dict.values())).keys())
    
    fig = go.Figure()
    
    for label, scores in scores_dict.items():
        fig.add_trace(go.Scatterpolar(
            r=[scores[cat] for cat in categories],
            theta=categories,
            fill='toself',
            name=label
        ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 1]
            )
        ),
        showlegend=True,
        title="Job Description Comparison",
        height=600
    )
    
    return fig

def create_comparison_dataframe(scores_dict):
    """Create a DataFrame comparing multiple job descriptions"""
    categories = list(next(iter(scores_dict.values())).keys())
    
    df_data = {
        'Category': categories,
    }
    
    # Add scores for each version
    for label, scores in scores_dict.items():
        df_data[label] = [f"{scores[cat]:.2%}" for cat in categories]
        
        # Calculate change from original if this isn't the original
        if label != 'Original':
            original_scores = scores_dict['Original']
            df_data[f'{label} Change'] = [
                f"{(scores[cat] - original_scores[cat])*100:+.2f}%" 
                for cat in categories
            ]
    
    return pd.DataFrame(df_data)

def init_session_state():
    """Initialize session state variables if they don't exist"""
    if 'username' not in st.session_state:
        st.session_state.username = 'Anonymous'
    
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
        
    if 'feedback_history' not in st.session_state:
        st.session_state.feedback_history = []
        
    if 'last_file' not in st.session_state:
        st.session_state.last_file = None
        
    if 'reload_flag' not in st.session_state:
        st.session_state.reload_flag = False
        
    if 'clear_feedback' not in st.session_state:
        st.session_state.clear_feedback = False
        
    if 'viewing_all_feedback' not in st.session_state:
        st.session_state.viewing_all_feedback = False
        
    if 'viewing_session_feedback' not in st.session_state:
        st.session_state.viewing_session_feedback = False

def get_or_create_logger():
    """Get existing logger from session state or create a new one"""
    # First check if we have a logger in session state
    if 'logger' in st.session_state:
        return st.session_state.logger
    
    # If we have a session_id, try to load that session
    if 'session_id' in st.session_state:
        try:
            # Try to load existing session by ID
            logger = JDOptimLogger.load_session(st.session_state.session_id)
            if logger:
                # Update username if it changed
                if logger.username != st.session_state.username:
                    logger.username = st.session_state.username
                    logger.current_state["username"] = st.session_state.username
                    logger._save_state()
                
                st.session_state.logger = logger
                return logger
        except Exception as e:
            # If loading fails, we'll create a new logger below
            print(f"Failed to load existing session: {e}")
    
    # Create a new logger with the current username
    logger = JDOptimLogger(username=st.session_state.username)
    st.session_state.session_id = logger.session_id
    st.session_state.logger = logger
    
    return logger

def render_username_input():
    """Render the username input in the sidebar"""
    st.sidebar.subheader("User Information")
    
    # Get current username from session state
    current_username = st.session_state.get('username', 'Anonymous')
    
    # Username input
    username = st.sidebar.text_input(
        "Your Name:",
        value=current_username,
        help="Enter your name to track your changes in the logs",
        key="username_input"
    )
    
    # Update username if changed
    if username != current_username:
        st.session_state.username = username
        
        # Update logger if it exists
        if 'logger' in st.session_state:
            st.session_state.logger.username = username
            st.session_state.logger.current_state["username"] = username
            st.session_state.logger._save_state()

def display_filtered_feedback_history():
    """Display feedback history with filtering options without creating new sessions"""
    # Get all available sessions
    sessions = JDOptimLogger.list_sessions()
    
    if not sessions:
        st.info("No previous sessions found with feedback")
        return
    
    # Create a list to store all feedback data
    all_feedback = []
    
    # Collect unique values for filters
    unique_users = set()
    unique_files = set()
    unique_dates = set()
    
    # Loop through each session to collect feedback
    for session_info in sessions:
        session_id = session_info["session_id"]
        try:
            # Load the session data directly from file without creating a new logger instance
            log_file = os.path.join("logs", f"jdoptim_session_{session_id}.json")
            if os.path.exists(log_file):
                with open(log_file, 'r') as f:
                    state = json.load(f)
                
                username = state.get("username", "Anonymous")
                unique_users.add(username)
                
                file_name = state.get("selected_file", "Unknown")
                unique_files.add(file_name)
                
                session_date = state.get("session_start_time", "Unknown")
                # Extract just the date part if it's a full timestamp
                if isinstance(session_date, str) and "T" in session_date:
                    session_date = session_date.split("T")[0]
                unique_dates.add(session_date)
                
                # Add each feedback item with metadata
                for i, feedback in enumerate(state.get("feedback_history", [])):
                    # Get timestamp for the feedback if available
                    feedback_time = "Unknown"
                    for action in state.get("actions", []):
                        if action.get("action") == "feedback" and action.get("index", -1) == i:
                            feedback_time = action.get("timestamp", "Unknown")
                            break
                    
                    # Handle different feedback formats (string or dict)
                    if isinstance(feedback, dict):
                        feedback_content = feedback.get("feedback", "")
                    else:
                        feedback_content = feedback
                    
                    all_feedback.append({
                        "Username": username,
                        "File": file_name,
                        "Session Date": session_date,
                        "Feedback Time": feedback_time,
                        "Feedback": feedback_content
                    })
        except Exception as e:
            print(f"Error reading session {session_id}: {str(e)}")
    
    if not all_feedback:
        st.info("No feedback found in any session")
        return
            
    # Convert to DataFrame
    feedback_df = pd.DataFrame(all_feedback)
    
    # Sort by most recent first if timestamps are available
    if "Feedback Time" in feedback_df.columns:
        try:
            # Parse timestamps where possible
            parsed_timestamps = []
            for timestamp in feedback_df["Feedback Time"]:
                try:
                    if isinstance(timestamp, str) and "T" in timestamp:
                        dt = datetime.datetime.fromisoformat(timestamp)
                        parsed_timestamps.append(dt)
                    else:
                        parsed_timestamps.append(datetime.datetime(1900, 1, 1))
                except:
                    parsed_timestamps.append(datetime.datetime(1900, 1, 1))
            
            feedback_df["Parsed Timestamp"] = parsed_timestamps
            feedback_df = feedback_df.sort_values("Parsed Timestamp", ascending=False)
        except:
            pass  # If sorting fails, just use the original order
    
    # Create filter sidebar with a container to avoid affecting main UI
    filter_container = st.container()
    
    with filter_container:
        st.subheader("Filter Feedback")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # User filter
            selected_users = st.multiselect(
                "Filter by User:",
                options=sorted(list(unique_users)),
                default=[]
            )
        
        with col2:
            # File filter
            selected_files = st.multiselect(
                "Filter by Job Description:",
                options=sorted(list(unique_files)),
                default=[]
            )
        
        with col3:
            # Date filter
            selected_dates = st.multiselect(
                "Filter by Date:",
                options=sorted(list(unique_dates), reverse=True),
                default=[]
            )
        
        # Text search
        search_text = st.text_input("Search in feedback:", "")
    
    # Apply filters
    filtered_df = feedback_df.copy()
    
    if selected_users:
        filtered_df = filtered_df[filtered_df["Username"].isin(selected_users)]
    
    if selected_files:
        filtered_df = filtered_df[filtered_df["File"].isin(selected_files)]
    
    if selected_dates:
        filtered_df = filtered_df[filtered_df["Session Date"].isin(selected_dates)]
    
    if search_text:
        filtered_df = filtered_df[filtered_df["Feedback"].str.contains(search_text, case=False, na=False)]
    
    # Show filter summary
    st.write(f"Showing {len(filtered_df)} of {len(feedback_df)} feedback items")
    
    # Display the filtered dataframe
    if not filtered_df.empty:
        # Format timestamps for display
        readable_timestamps = []
        for timestamp in filtered_df["Feedback Time"]:
            try:
                if isinstance(timestamp, str) and "T" in timestamp:
                    dt = datetime.datetime.fromisoformat(timestamp)
                    readable_timestamps.append(dt.strftime("%Y-%m-%d %H:%M:%S"))
                else:
                    readable_timestamps.append(str(timestamp))
            except:
                readable_timestamps.append(str(timestamp))
        
        filtered_df["Formatted Time"] = readable_timestamps
        
        # Display dataframe
        st.dataframe(
            filtered_df,
            use_container_width=True,
            column_config={
                "Username": st.column_config.TextColumn("User"),
                "File": st.column_config.TextColumn("Job Description"),
                "Formatted Time": st.column_config.TextColumn("Time"),
                "Feedback": st.column_config.TextColumn("Feedback Content", width="large"),
            },
            hide_index=True
        )
    else:
        st.info("No feedback matches the selected filters")
    
    # Option to export filtered results
    if not filtered_df.empty:
        csv = filtered_df.to_csv(index=False)
        st.download_button(
            label="📥 Export Filtered Feedback",
            data=csv,
            file_name=f"feedback_export_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )

def cleanup_old_logs(max_sessions=50, older_than_days=30):
    """Clean up old log files"""
    log_dir = "logs"
    if not os.path.exists(log_dir):
        return
    
    # Get all log files with their modification times
    files_with_time = []
    for filename in os.listdir(log_dir):
        if filename.startswith("jdoptim_session_") and filename.endswith(".json"):
            file_path = os.path.join(log_dir, filename)
            mod_time = os.path.getmtime(file_path)
            files_with_time.append((file_path, mod_time))
    
    # Sort by modification time (newest first)
    files_with_time.sort(key=lambda x: x[1], reverse=True)
    
    # Keep recent sessions (max_sessions)
    keep_files = [file_path for file_path, _ in files_with_time[:max_sessions]]
    
    # Delete older files
    current_time = time.time()
    cutoff_time = current_time - (older_than_days * 24 * 60 * 60)
    
    deleted_count = 0
    for file_path, mod_time in files_with_time:
        # Only delete if it's not in the keep_files list and older than cutoff
        if file_path not in keep_files and mod_time < cutoff_time:
            try:
                os.remove(file_path)
                deleted_count += 1
            except:
                pass
    
    return deleted_count

def render_sidebar_with_history(logger):
    """Render the sidebar with session history and controls"""
    # First render the username input
    render_username_input()
    
    st.sidebar.title("Session History")
    
    # Get all available sessions
    sessions = JDOptimLogger.list_sessions()
    
    # Display current session stats
    st.sidebar.subheader("Current Session")
    
    st.sidebar.caption(f"User: {logger.username}")
    st.sidebar.caption(f"Session ID: {logger.session_id[:8]}...")
    
    if len(logger.current_state["actions"]) > 0:
        actions_summary = {}
        for action in logger.current_state["actions"]:
            action_type = action["action"]
            if action_type in actions_summary:
                actions_summary[action_type] += 1
            else:
                actions_summary[action_type] = 1
        
        # Create a formatted string of actions
        actions_text = ", ".join([f"{count} {action}" for action, count in actions_summary.items()])
        st.sidebar.caption(f"Actions: {actions_text}")
        
        if logger.current_state["selected_file"]:
            st.sidebar.caption(f"File: {logger.current_state['selected_file']}")
        
        if logger.current_state["feedback_history"]:
            st.sidebar.caption(f"Feedback items: {len(logger.current_state['feedback_history'])}")
    
    # Session controls
    st.sidebar.button("🔄 Start New Session", on_click=start_new_session)
    
    # Display previous sessions if any
    if sessions:
        st.sidebar.subheader(f"Previous Sessions ({len(sessions)})")
        
        # Show the 5 most recent sessions
        for i, session in enumerate(sessions[:5]):
            session_id = session["session_id"]
            file_name = session.get("file_processed", "Unknown")
            user = session.get("username", "Anonymous")
            start_time = session.get("start_time", "Unknown")
            
            session_display = f"{user} - {file_name}"
            if len(session_display) > 30:
                session_display = session_display[:27] + "..."
                
            with st.sidebar.expander(f"{session_display} ({start_time})"):
                st.caption(f"User: {user}")
                st.caption(f"Session ID: {session_id[:8]}...")
                
                # Add button to load this session
                if st.button("Load Session", key=f"load_session_{i}"):
                    load_previous_session(session_id)

def start_new_session():
    """Start a new session"""
    # Clean up session state
    for key in ['logger', 'session_id', 'enhanced_versions', 'original_jd', 
                'feedback_history', 'current_enhanced_version', 'last_file']:
        if key in st.session_state:
            del st.session_state[key]
    
    # Flag for reload
    st.session_state.reload_flag = True

def load_previous_session(session_id):
    """Load a previous session by ID"""
    try:
        logger = JDOptimLogger.load_session(session_id)
        if logger:
            # Update session state
            st.session_state.logger = logger
            st.session_state.session_id = logger.session_id
            st.session_state.username = logger.username
            
            # Restore feedback history and other state
            st.session_state.feedback_history = logger.current_state["feedback_history"]
            st.session_state.current_enhanced_version = logger.current_state["current_enhanced_version"]
            st.session_state.last_file = logger.current_state["selected_file"]
            
            # Set reload flag
            st.session_state.reload_flag = True
            
            st.success(f"Loaded session successfully")
            st.rerun()
        else:
            st.error("Failed to load session")
    except Exception as e:
        st.error(f"Error loading session: {str(e)}")

def main():
    # Set page config
    st.set_page_config(
        page_title="Job Description Enhancer",
        page_icon="💼",
        layout="wide"
    )
    
    # Initialize session state
    init_session_state()
    
    # Get or create logger
    logger = get_or_create_logger()
    
    st.title("💼 Job Description Enhancer")
    st.markdown("Select a job description file, choose from enhanced versions, and provide feedback for final enhancement")

    # Initialize the analyzer and agent
    analyzer = JobDescriptionAnalyzer()
    agent = JobDescriptionAgent(model_id="anthropic.claude-3-haiku-20240307-v1:0")
    
    # Render the sidebar with session history
    render_sidebar_with_history(logger)
    
    # Clean up old logs in the background
    cleanup_old_logs()

    # File selection
    jd_directory = os.path.join(os.getcwd(), "JDs")
    try:
        files = [f for f in os.listdir(jd_directory) if f.endswith(('.txt', '.docx'))]
    except FileNotFoundError:
        # If directory not found, allow direct file upload
        st.warning("Directory 'JDs' not found. You can upload a job description file directly.")
        uploaded_file = st.file_uploader("Upload Job Description File", type=['txt', 'docx'])
        if uploaded_file:
            if uploaded_file.name.endswith('.txt'):
                content = uploaded_file.getvalue().decode('utf-8')
            else:  # .docx
                doc = Document(uploaded_file)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            st.session_state.original_jd = content
            selected_file = uploaded_file.name
            
            # Log file selection (only if changed)
            if logger.current_state["selected_file"] != selected_file:
                logger.log_file_selection(selected_file, content)
        else:
            st.error("Please either upload a file or create a 'JDs' folder in the application directory.")
            return
    else:
        # Create columns for file selection and session information
        file_col, session_col = st.columns([2, 1])

        with file_col:
            selected_file = st.selectbox(
                "Select Job Description File",
                files,
                help="Choose a job description file to enhance",
                key="file_selector"
            )

        with session_col:
            # Display session information
            st.info(f"Session: {logger.session_id[:8]}...")
            st.caption(f"User: {logger.username}")
            if len(logger.current_state["actions"]) > 0:
                st.caption(f"Actions: {len(logger.current_state['actions'])}")
                if logger.current_state["selected_file"]:
                    st.caption(f"Working on: {logger.current_state['selected_file']}")

        # Reset state when file changes
        if st.session_state.last_file != selected_file:
            st.session_state.last_file = selected_file
            if 'enhanced_versions' in st.session_state:
                del st.session_state.enhanced_versions
            if 'original_jd' in st.session_state:
                del st.session_state.original_jd
            st.session_state.reload_flag = True
            
        if selected_file:
            file_path = os.path.join(jd_directory, selected_file)
            try:
                # Read the job description
                if 'original_jd' not in st.session_state:
                    st.session_state.original_jd = read_job_description(file_path)
                    # Log file selection (only if changed)
                    if logger.current_state["selected_file"] != selected_file:
                        logger.log_file_selection(selected_file, st.session_state.original_jd)
            except Exception as e:
                st.error(f"Error reading file: {str(e)}")
                return

    # From here, the rest of the app continues with either the uploaded file or selected file
    original_jd = st.session_state.original_jd
            
    # Display original JD
    st.subheader("Original Job Description")
    st.text_area("Original Content", original_jd, height=200, disabled=True, key="original_jd_display")

    # Generate button
    generate_btn = st.button("✨ Generate Enhanced Versions", type="primary", key="generate_btn")
    
    # Handle generating enhanced versions
    if generate_btn or ('enhanced_versions' not in st.session_state and st.session_state.reload_flag):
        st.session_state.reload_flag = False
        with st.spinner("Generating enhanced versions..."):
            versions = agent.generate_initial_descriptions(original_jd)
            # Ensure we have 3 versions
            while len(versions) < 3:
                versions.append(f"Enhanced Version {len(versions)+1}:\n{original_jd}")
            st.session_state.enhanced_versions = versions
            # Log versions generation (only once)
            logger.log_versions_generated(versions)
            # Use rerun to refresh the UI
            st.rerun()
    
    # If enhanced versions are available, display them
    if 'enhanced_versions' in st.session_state and len(st.session_state.enhanced_versions) >= 3:
        # Analyze all versions
        original_scores = analyzer.analyze_text(original_jd)
        intermediate_scores = {
            f'Version {i+1}': analyzer.analyze_text(version)
            for i, version in enumerate(st.session_state.enhanced_versions)
        }
        
        # Combine all scores for comparison
        all_scores = {'Original': original_scores, **intermediate_scores}

        # Display enhanced versions and their analysis
        st.subheader("Enhanced Versions Comparison")
        
        # Create tabs for content and analysis
        content_tab, analysis_tab = st.tabs(["Content", "Analysis"])
        
        with content_tab:
            enhanced_versions_tabs = st.tabs(["Version 1", "Version 2", "Version 3"])
            for idx, (tab, version) in enumerate(zip(enhanced_versions_tabs, st.session_state.enhanced_versions)):
                with tab:
                    st.text_area(
                        f"Enhanced Version {idx + 1}",
                        version,
                        height=300,
                        disabled=True,
                        key=f"enhanced_version_{idx}"
                    )

        with analysis_tab:
            col1, col2 = st.columns([1, 1])
            
            with col1:
                radar_chart = create_multi_radar_chart(all_scores)
                st.plotly_chart(radar_chart, use_container_width=True, key="intermediate_radar")
            
            with col2:
                comparison_df = create_comparison_dataframe(all_scores)
                st.dataframe(
                    comparison_df,
                    height=500,
                    use_container_width=True,
                    hide_index=True,
                    key="intermediate_comparison"
                )

        # Version selection
        st.subheader("Select Version for Final Enhancement")
        selected_version = st.radio(
            "Choose the version you'd like to use as a base:",
            ["Version 1", "Version 2", "Version 3"],
            help="Select the version that best matches your needs for further enhancement",
            key="version_selector"
        )
        
        # Get selected version index
        selected_index = int(selected_version[-1]) - 1
        
        # User feedback
        st.subheader("Provide Additional Feedback")
        
        # Display feedback history
        if logger.current_state["feedback_history"]:
            st.write("Previous Feedback:")
            for i, feedback in enumerate(logger.current_state["feedback_history"], 1):
                if isinstance(feedback, dict):
                    feedback_text = feedback.get("feedback", "")
                else:
                    feedback_text = feedback
                st.text(f"{i}. {feedback_text}")
        
        # Handle feedback clearing mechanism
        if st.session_state.get('clear_feedback', False):
            # Reset the flag
            st.session_state.clear_feedback = False
            # The form will be empty on this rerun because we're not restoring any value
        
        user_feedback = st.text_area(
            "Enter your feedback for further enhancement",
            height=100,
            placeholder="E.g., 'Add more emphasis on leadership skills', 'Include cloud technologies', etc.",
            key="user_feedback"
        )
        
        # Add a button to just save feedback without generating final version
        if st.button("➕ Add Feedback Without Generating Final Version", type="secondary", key="add_feedback_only"):
            if user_feedback.strip():
                # Log the feedback
                logger.log_feedback(user_feedback)
                # Use a different session state variable to trigger clearing the text area on next rerun
                st.session_state.clear_feedback = True
                st.success("Feedback added successfully! You can add more feedback or generate the final version when ready.")
                st.rerun()
            else:
                st.warning("Please enter some feedback first.")
        
        # Feedback history buttons
        st.markdown("---")
        feedback_col1, feedback_col2 = st.columns([1, 1])
        
        with feedback_col1:
            if st.button("👁️ View All Previous Feedback", type="secondary", key="view_all_feedback"):
                # Don't generate a new session for viewing feedback
                st.session_state['viewing_all_feedback'] = True
        
        with feedback_col2:
            if st.button("📋 View Current Session Feedback", type="secondary", key="view_session_feedback"):
                # Don't generate a new session for viewing session feedback
                st.session_state['viewing_session_feedback'] = True
        
        # Display feedback if requested
        if st.session_state.get('viewing_all_feedback', False):
            st.subheader("Feedback History from All Sessions")
            # Use the direct file reading method instead of loading sessions
            display_filtered_feedback_history()
            # Reset viewing flag after displaying
            st.session_state['viewing_all_feedback'] = False
            
        if st.session_state.get('viewing_session_feedback', False):
            st.subheader("Current Session Feedback History")
            # Simply display the feedback from the current logger
            if logger.current_state["feedback_history"]:
                for i, feedback in enumerate(logger.current_state["feedback_history"], 1):
                    if isinstance(feedback, dict):
                        feedback_text = feedback.get("feedback", "")
                    else:
                        feedback_text = feedback
                    st.text(f"{i}. {feedback_text}")
            else:
                st.info("No feedback in current session")
            # Reset viewing flag after displaying
            st.session_state['viewing_session_feedback'] = False
            
        # Final enhancement process
        final_btn_label = f"🚀 Generate Final Version Using All Feedback ({len(logger.current_state['feedback_history'])} items)"
        if st.button(final_btn_label, type="primary", key="generate_final"):
            try:
                with st.spinner("Creating final enhanced version..."):
                    # Log version selection
                    logger.log_version_selection(selected_index)
                    
                    # Log new feedback if provided and not already added
                    if user_feedback.strip():
                        logger.log_feedback(user_feedback)
                        # Use a different session state variable to trigger clearing the text area on next rerun
                        st.session_state.clear_feedback = True
                    
                    # Use the current enhanced version if it exists, otherwise use selected version
                    base_description = (logger.current_state["current_enhanced_version"] or 
                                      st.session_state.enhanced_versions[selected_index])
                    
                    # Create list of feedback (convert dicts to strings if needed)
                    feedback_list = []
                    for item in logger.current_state["feedback_history"]:
                        if isinstance(item, dict):
                            feedback_list.append(item.get("feedback", ""))
                        else:
                            feedback_list.append(item)
                    
                    # Generate final description using accumulated feedback
                    final_description = agent.generate_final_description(
                        base_description,
                        feedback_list
                    )
                    
                    # Log the new enhanced version
                    logger.log_enhanced_version(final_description, is_final=True)
                    
                    # Store the new enhanced version in session state
                    st.session_state.current_enhanced_version = final_description
                    
                    # Add final version to scores dictionary
                    final_scores = analyzer.analyze_text(final_description)
                    all_scores['Final'] = final_scores
                    
                    # Display final version
                    st.subheader("Final Enhanced Job Description")
                    st.text_area(
                        "Final Content",
                        final_description,
                        height=400,
                        key="final_description"
                    )
                    
                    # Create comparison section
                    st.subheader("Final Comparison Analysis")
                    final_col1, final_col2 = st.columns([1, 1])
                    
                    with final_col1:
                        final_radar = create_multi_radar_chart(all_scores)
                        st.plotly_chart(final_radar, use_container_width=True, key="final_radar")
                    
                    with final_col2:
                        final_comparison_df = create_comparison_dataframe(all_scores)
                        st.dataframe(
                            final_comparison_df,
                            height=500,
                            use_container_width=True,
                            hide_index=True,
                            key="final_comparison"
                        )
                    
                    # Download options
                    st.markdown("---")
                    download_col1, download_col2, report_col = st.columns(3)
                    
                    with download_col1:
                        st.download_button(
                            label="📥 Download as TXT",
                            data=final_description,
                            file_name=f"enhanced_jd_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                            key="download_txt"
                        )
                        # Log download action
                        logger.log_download("txt", f"enhanced_jd_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
                    
                    with download_col2:
                        if st.button("📥 Download as DOCX", key="download_docx"):
                            docx_filename = f"enhanced_jd_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                            save_enhanced_jd(final_description, docx_filename, 'docx')
                            st.success(f"Saved as {docx_filename}")
                            # Log download action
                            logger.log_download("docx", docx_filename)
                    
                    with report_col:
                        if st.button("📊 Generate Session Report", key="generate_report"):
                            report_file = logger.export_session_report()
                            st.success(f"Session report generated: {os.path.basename(report_file)}")
                            with open(report_file, 'r') as f:
                                report_content = f.read()
                            st.download_button(
                                label="📥 Download Report",
                                data=report_content,
                                file_name=os.path.basename(report_file),
                                mime="text/plain",
                                key="download_report"
                            )
                            
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
                st.error("Please try again or contact support if the problem persists.")
            
            # Add button to clear feedback history
            if st.button("Clear Feedback History", key="clear_feedback"):
                # Reset feedback in the logger
                logger.current_state["feedback_history"] = []
                logger.current_state["current_enhanced_version"] = None
                logger._save_state()
                # Also reset in session state for UI consistency
                st.session_state.feedback_history = []
                st.session_state.current_enhanced_version = None
                st.rerun()

    # Footer
    st.markdown("---")
    st.markdown("Made by Apexon |  Use the feedback section to customize the enhancement | Uses logging to track session state and enable history review")

if __name__ == "__main__":
    main()