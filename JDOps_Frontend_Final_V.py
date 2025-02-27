import streamlit as st
import boto3
import json
import pandas as pd
import os
from docx import Document
import re
from collections import Counter
import numpy as np
import plotly.graph_objects as go
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import nltk

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    
def read_job_description(file_path):
    """Read job description from either .txt or .docx file"""
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format")

# Function to save enhanced JD
def save_enhanced_jd(content, filename, format_type):
    if format_type == 'docx':
        doc = Document()
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(filename)
        return True
    return False

class JobDescriptionAnalyzer:
    def __init__(self):
        self.categories = {
            'Technical Skills': ['python', 'java', 'sql', 'aws', 'cloud', 'docker', 'kubernetes', 'api', 'database', 
                               'git', 'linux', 'agile', 'devops', 'ml', 'ai', 'analytics'],
            'Soft Skills': ['communication', 'leadership', 'teamwork', 'collaboration', 'problem-solving', 'analytical', 
                           'initiative', 'organizational', 'time management', 'interpersonal'],
            'Experience Level': ['year', 'years', 'senior', 'junior', 'mid-level', 'lead', 'manager', 'experience'],
            'Education': ['degree', 'bachelor', 'master', 'phd', 'certification', 'education'],
            'Tools & Technologies': ['jira', 'confluence', 'slack', 'github', 'gitlab', 'azure', 'jenkins', 'terraform'],
            'Domain Knowledge': ['finance', 'healthcare', 'retail', 'banking', 'insurance', 'technology', 'manufacturing']
        }

    def analyze_text(self, text):
        text = text.lower()
        scores = {}
        
        for category, keywords in self.categories.items():
            category_score = 0
            for keyword in keywords:
                count = len(re.findall(r'\b' + keyword + r'\b', text))
                category_score += count
            max_possible = len(keywords)
            scores[category] = min(category_score / max_possible, 1.0)
            
        return scores

class JobDescriptionAgent:
    def __init__(self, model_id, max_tokens=1500, temperature=0.7):
        self.model_id = model_id
        self.max_tokens = max_tokens
        self.temperature = temperature
        self.client = boto3.client(
            service_name='bedrock-runtime',
            aws_access_key_id='<key>',
            aws_secret_access_key='<key>',
            region_name='us-east-1',
        )

    def generate_initial_descriptions(self, job_description):
        prompt = (
            "Using the job description provided, create three distinct enhanced versions. Each should expand on these sections:\n\n"
            "1. Brief overview of the role\n"
            "2. Key responsibilities\n"
            "3. Required skills\n"
            "4. Preferred skills\n"
            "5. Required experience\n"
            "6. Preferred experience\n"
            "7. Required tools\n\n"
            "Format your response exactly as follows:\n"
            "VERSION 1:\n[First enhanced version]\n\n"
            "VERSION 2:\n[Second enhanced version]\n\n"
            "VERSION 3:\n[Third enhanced version]\n\n"
            "Make each version unique while maintaining accuracy and relevance.\n\n"
            f"Original Job Description:\n{job_description}"
        )

        native_request = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": self.max_tokens,
            "temperature": self.temperature,
            "messages": [{"role": "user", "content": prompt}],
        }
        
        response = self.client.invoke_model(
            modelId=self.model_id,
            body=json.dumps(native_request),
            contentType="application/json",
        )
        response_body = response['body'].read().decode("utf-8")
        model_response = json.loads(response_body)

        if "content" in model_response and isinstance(model_response["content"], list):
            full_text = model_response["content"][0]["text"].strip()
            
            # More robust splitting pattern
            parts = re.split(r'VERSION \d+:', full_text)
            if len(parts) >= 4:  # The first part is empty or intro text
                descriptions = [part.strip() for part in parts[1:4]]
                return descriptions
            else:
                # Fallback parsing method
                descriptions = []
                version_pattern = re.compile(r'VERSION (\d+):(.*?)(?=VERSION \d+:|$)', re.DOTALL)
                matches = version_pattern.findall(full_text)
                for _, content in matches[:3]:
                    descriptions.append(content.strip())
                
                if len(descriptions) == 3:
                    return descriptions
        
        # If we failed to parse properly, generate simpler versions
        return [
            f"Enhanced Version 1 of the job description:\n{job_description}",
            f"Enhanced Version 2 of the job description:\n{job_description}",
            f"Enhanced Version 3 of the job description:\n{job_description}"
        ]

    def generate_final_description(self, selected_description, feedback_history):
        """
        Generate enhanced description incorporating feedback history
        
        Args:
            selected_description (str): The base description to enhance
            feedback_history (list): List of previous feedback items
        """
        # Construct prompt with feedback history
        feedback_context = "\n".join([
            f"Previous Feedback {i+1}: {feedback}" 
            for i, feedback in enumerate(feedback_history[:-1])
        ])
        
        current_feedback = feedback_history[-1] if feedback_history else ""
        
        prompt = (
            "Using the selected job description and incorporating all previous feedback plus new feedback, "
            "create an enhanced version.\n\n"
            f"Selected Job Description:\n{selected_description}\n\n"
        )
        
        if feedback_context:
            prompt += f"Previous Feedback Applied:\n{feedback_context}\n\n"
        
        prompt += (
            f"New Feedback to Implement:\n{current_feedback}\n\n"
            "Enhance while maintaining this structure:\n"
            "1. Brief overview of the role\n"
            "2. Key responsibilities\n"
            "3. Required skills\n"
            "4. Preferred skills\n"
            "5. Required experience\n"
            "6. Preferred experience\n"
            "7. Required tools\n\n"
            "Important: Ensure all previous feedback improvements are preserved while implementing the new feedback."
        )

        native_request = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": self.max_tokens,
            "temperature": self.temperature,
            "messages": [{"role": "user", "content": prompt}],
        }
        
        response = self.client.invoke_model(
            modelId=self.model_id,
            body=json.dumps(native_request),
            contentType="application/json",
        )
        response_body = response['body'].read().decode("utf-8")
        model_response = json.loads(response_body)

        if "content" in model_response and isinstance(model_response["content"], list):
            return model_response["content"][0]["text"].strip()
        else:
            raise ValueError("Unexpected response format")

def create_multi_radar_chart(scores_dict):
    """Create a radar chart comparing multiple job descriptions"""
    categories = list(next(iter(scores_dict.values())).keys())
    
    fig = go.Figure()
    
    for label, scores in scores_dict.items():
        fig.add_trace(go.Scatterpolar(
            r=[scores[cat] for cat in categories],
            theta=categories,
            fill='toself',
            name=label
        ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 1]
            )
        ),
        showlegend=True,
        title="Job Description Comparison",
        height=600
    )
    
    return fig

def create_comparison_dataframe(scores_dict):
    """Create a DataFrame comparing multiple job descriptions"""
    categories = list(next(iter(scores_dict.values())).keys())
    
    df_data = {
        'Category': categories,
    }
    
    # Add scores for each version
    for label, scores in scores_dict.items():
        df_data[label] = [f"{scores[cat]:.2%}" for cat in categories]
        
        # Calculate change from original if this isn't the original
        if label != 'Original':
            original_scores = scores_dict['Original']
            df_data[f'{label} Change'] = [
                f"{(scores[cat] - original_scores[cat])*100:+.2f}%" 
                for cat in categories
            ]
    
    return pd.DataFrame(df_data)

# Set page config at the very beginning
st.set_page_config(
    page_title="Job Description Enhancer",
    page_icon="💼",
    layout="wide"
)


def main():
    st.title("💼 Job Description Enhancer")
    st.markdown("Select a job description file, choose from enhanced versions, and provide feedback for final enhancement")

    # Initialize the analyzer and agent
    analyzer = JobDescriptionAnalyzer()
    agent = JobDescriptionAgent(model_id="anthropic.claude-3-haiku-20240307-v1:0")

    # Initialize session state
    if 'last_file' not in st.session_state:
        st.session_state.last_file = None
    if 'reload_flag' not in st.session_state:
        st.session_state.reload_flag = False
        
    if 'feedback_history' not in st.session_state:
        st.session_state.feedback_history = []
    if 'current_enhanced_version' not in st.session_state:
        st.session_state.current_enhanced_version = None
    # File selection
    jd_directory = "/Users/raghav/Desktop/Apexon/JD Optimization/JDs"  # Update this path as needed
    try:
        files = [f for f in os.listdir(jd_directory) if f.endswith(('.txt', '.docx'))]
    except FileNotFoundError:
        st.error(f"Directory '{jd_directory}' not found. Please create a 'JDs' folder and add your job description files.")
        return

    if not files:
        st.warning("No .txt or .docx files found in the JDs folder. Please add some files and refresh.")
        return

    selected_file = st.selectbox(
        "Select Job Description File",
        files,
        help="Choose a job description file to enhance",
        key="file_selector"
    )

    # Reset session state when file changes
    if st.session_state.last_file != selected_file:
        st.session_state.last_file = selected_file
        if 'enhanced_versions' in st.session_state:
            del st.session_state.enhanced_versions
        if 'original_jd' in st.session_state:
            del st.session_state.original_jd
        st.session_state.reload_flag = True

    if selected_file:
        file_path = os.path.join(jd_directory, selected_file)
        try:
            # Read the job description
            if 'original_jd' not in st.session_state:
                st.session_state.original_jd = read_job_description(file_path)
            
            original_jd = st.session_state.original_jd
            
            # Display original JD
            st.subheader("Original Job Description")
            st.text_area("Original Content", original_jd, height=200, disabled=True, key="original_jd_display")

            # Generate button
            generate_btn = st.button("✨ Generate Enhanced Versions", type="primary", key="generate_btn")
            
            # Handle generating enhanced versions
            if generate_btn or ('enhanced_versions' not in st.session_state and st.session_state.reload_flag):
                st.session_state.reload_flag = False
                with st.spinner("Generating enhanced versions..."):
                    versions = agent.generate_initial_descriptions(original_jd)
                    # Ensure we have 3 versions
                    while len(versions) < 3:
                        versions.append(f"Enhanced Version {len(versions)+1}:\n{original_jd}")
                    st.session_state.enhanced_versions = versions
                    # Use st.rerun() instead of experimental_rerun
                    st.rerun()

            # If enhanced versions are available, display them
            if 'enhanced_versions' in st.session_state and len(st.session_state.enhanced_versions) >= 3:
                # Analyze all versions
                original_scores = analyzer.analyze_text(original_jd)
                intermediate_scores = {
                    f'Version {i+1}': analyzer.analyze_text(version)
                    for i, version in enumerate(st.session_state.enhanced_versions)
                }
                
                # Combine all scores for comparison
                all_scores = {'Original': original_scores, **intermediate_scores}

                # Display enhanced versions and their analysis
                st.subheader("Enhanced Versions Comparison")
                
                # Create tabs for content and analysis
                content_tab, analysis_tab = st.tabs(["Content", "Analysis"])
                
                with content_tab:
                    enhanced_versions_tabs = st.tabs(["Version 1", "Version 2", "Version 3"])
                    for idx, (tab, version) in enumerate(zip(enhanced_versions_tabs, st.session_state.enhanced_versions)):
                        with tab:
                            st.text_area(
                                f"Enhanced Version {idx + 1}",
                                version,
                                height=300,
                                disabled=True,
                                key=f"enhanced_version_{idx}"
                            )

                with analysis_tab:
                    col1, col2 = st.columns([1, 1])
                    
                    with col1:
                        radar_chart = create_multi_radar_chart(all_scores)
                        st.plotly_chart(radar_chart, use_container_width=True, key="intermediate_radar")
                    
                    with col2:
                        comparison_df = create_comparison_dataframe(all_scores)
                        st.dataframe(
                            comparison_df,
                            height=500,
                            use_container_width=True,
                            hide_index=True,
                            key="intermediate_comparison"
                        )

                # Version selection
                st.subheader("Select Version for Final Enhancement")
                selected_version = st.radio(
                    "Choose the version you'd like to use as a base:",
                    ["Version 1", "Version 2", "Version 3"],
                    help="Select the version that best matches your needs for further enhancement",
                    key="version_selector"
                )
                
                # User feedback
                st.subheader("Provide Additional Feedback")
                if st.session_state.feedback_history:
                    st.write("Previous Feedback:")
                    for i, feedback in enumerate(st.session_state.feedback_history, 1):
                        st.text(f"{i}. {feedback}")
                
                user_feedback = st.text_area(
                    "Enter your feedback for further enhancement",
                    height=100,
                    placeholder="E.g., 'Add more emphasis on leadership skills', 'Include cloud technologies', etc.",
                    key="user_feedback"
                )
            
                # Final enhancement process
                if st.button("🚀 Generate Final Version", type="primary", key="generate_final"):
                    try:
                        with st.spinner("Creating final enhanced version..."):
                            # Get selected version index
                            selected_index = int(selected_version[-1]) - 1
                            
                            # Add new feedback to history
                            if user_feedback.strip():
                                st.session_state.feedback_history.append(user_feedback)
                            
                            # Use the current enhanced version if it exists, otherwise use selected version
                            base_description = (st.session_state.current_enhanced_version or 
                                              st.session_state.enhanced_versions[selected_index])
                            
                            # Generate final description using accumulated feedback
                            final_description = agent.generate_final_description(
                                base_description,
                                st.session_state.feedback_history
                            )
                            
                            # Store the new enhanced version
                            st.session_state.current_enhanced_version = final_description
                            
                            # Add final version to scores dictionary
                            final_scores = analyzer.analyze_text(final_description)
                            all_scores['Final'] = final_scores
                            
                            # Display final version
                            st.subheader("Final Enhanced Job Description")
                            st.text_area(
                                "Final Content",
                                final_description,
                                height=400,
                                key="final_description"
                            )
                            
                            # Create comparison section
                            st.subheader("Final Comparison Analysis")
                            final_col1, final_col2 = st.columns([1, 1])
                            
                            with final_col1:
                                final_radar = create_multi_radar_chart(all_scores)
                                st.plotly_chart(final_radar, use_container_width=True, key="final_radar")
                            
                            with final_col2:
                                final_comparison_df = create_comparison_dataframe(all_scores)
                                st.dataframe(
                                    final_comparison_df,
                                    height=500,
                                    use_container_width=True,
                                    hide_index=True,
                                    key="final_comparison"
                                )
                            
                            # Download options
                            st.markdown("---")
                            download_col1, download_col2 = st.columns(2)
                            
                            with download_col1:
                                st.download_button(
                                    label="📥 Download as TXT",
                                    data=final_description,
                                    file_name="enhanced_jd.txt",
                                    mime="text/plain",
                                    key="download_txt"
                                )
                            
                            with download_col2:
                                if st.button("📥 Download as DOCX", key="download_docx"):
                                    docx_filename = "enhanced_jd.docx"
                                    save_enhanced_jd(final_description, docx_filename, 'docx')
                                    st.success(f"Saved as {docx_filename}")
                                    
                    except Exception as e:
                        st.error(f"An error occurred: {str(e)}")
                        st.error("Please try again or contact support if the problem persists.")
                
                # Add button to clear feedback history
                if st.button("Clear Feedback History", key="clear_feedback"):
                    st.session_state.feedback_history = []
                    st.session_state.current_enhanced_version = None
                    st.rerun()

        except Exception as e:
            st.error(f"Error reading file: {str(e)}")
            return

    # Footer
    st.markdown("---")
    st.markdown("Made by Apexon | Use the feedback section to customize the enhancement")

if __name__ == "__main__":
    main()